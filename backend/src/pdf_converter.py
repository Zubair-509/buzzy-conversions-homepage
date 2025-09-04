#!/usr/bin/env python3
"""
Enhanced PDF to Word Converter with Multiple Conversion Modes
Supports Auto, Fast, Accurate, and Hybrid conversion modes with advanced formatting preservation
"""

import os
import io
import uuid
import tempfile
import numpy as np
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
import cv2
from PIL import Image
import logging

# Core libraries
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Mode-specific libraries
from pdf2docx import Converter  # Fast mode
import pytesseract  # OCR for Auto mode
from pdf2image import convert_from_path  # Hybrid mode

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class EnhancedPDFConverter:
    """Enhanced PDF to Word converter with multiple conversion modes"""
    
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
        self.image_counter = 0
        
    def detect_pdf_type(self, pdf_path: str) -> str:
        """Detect if PDF is scanned (image-based) or digital (text-based)"""
        try:
            pdf_doc = fitz.open(pdf_path)
            total_text_length = 0
            total_images = 0
            pages_checked = min(3, pdf_doc.page_count)  # Check first 3 pages
            
            for page_num in range(pages_checked):
                page = pdf_doc[page_num]
                text = page.get_text().strip()
                total_text_length += len(text)
                total_images += len(page.get_images())
            
            pdf_doc.close()
            
            # Heuristic: if very little text and many images, likely scanned
            if total_text_length < 100 and total_images > 0:
                return "scanned"
            elif total_text_length > 500:
                return "digital"
            else:
                return "mixed"
                
        except Exception as e:
            logger.warning(f"Error detecting PDF type: {e}")
            return "digital"  # Default assumption
    
    def convert_fast_mode(self, pdf_path: str, output_path: str) -> Dict[str, Any]:
        """Fast conversion using pdf2docx with bullet point enhancement"""
        try:
            logger.info("Using Fast mode (pdf2docx)")
            
            cv = Converter(pdf_path)
            cv.convert(output_path)
            cv.close()
            
            # Post-process the document to enhance bullet points
            self._enhance_fast_mode_bullets(output_path)
            
            # Get page count for reporting
            pdf_doc = fitz.open(pdf_path)
            pages_processed = pdf_doc.page_count
            pdf_doc.close()
            
            return {
                "success": True,
                "message": "PDF converted using Fast mode (pdf2docx) with bullet enhancement",
                "pages_processed": pages_processed,
                "mode": "fast",
                "output_file": output_path
            }
            
        except Exception as e:
            logger.error(f"Fast mode conversion failed: {e}")
            return {
                "success": False,
                "error": f"Fast mode conversion failed: {str(e)}",
                "mode": "fast"
            }
    
    def convert_accurate_mode(self, pdf_path: str, output_path: str) -> Dict[str, Any]:
        """Accurate conversion using PyMuPDF + python-docx - preserves complex formatting"""
        pdf_doc = None
        try:
            logger.info("Using Accurate mode (PyMuPDF + python-docx)")
            
            # Open PDF
            pdf_doc = fitz.open(pdf_path)
            if pdf_doc.page_count == 0:
                return {
                    "success": False,
                    "error": "PDF file is empty or corrupted",
                    "mode": "accurate"
                }
            
            # Create Word document
            word_doc = Document()
            pages_processed = 0
            
            # Process each page
            for page_num in range(pdf_doc.page_count):
                try:
                    page = pdf_doc[page_num]
                    
                    # Add page break if not first page
                    if page_num > 0:
                        word_doc.add_page_break()
                    
                    # Extract text with formatting
                    self._extract_formatted_text(page, word_doc)
                    
                    # Extract images
                    self._extract_images_accurate(page, word_doc, pdf_doc)
                    
                    # Extract tables
                    self._extract_tables_accurate(page, word_doc)
                    
                    pages_processed += 1
                    
                except Exception as e:
                    logger.warning(f"Error processing page {page_num + 1}: {e}")
                    continue
            
            # Save document
            word_doc.save(output_path)
            pdf_doc.close()
            
            # Cleanup
            self._cleanup_temp_files()
            
            return {
                "success": True,
                "message": "PDF converted using Accurate mode (PyMuPDF + python-docx)",
                "pages_processed": pages_processed,
                "mode": "accurate",
                "output_file": output_path
            }
            
        except Exception as e:
            if pdf_doc:
                pdf_doc.close()
            logger.error(f"Accurate mode conversion failed: {e}")
            return {
                "success": False,
                "error": f"Accurate mode conversion failed: {str(e)}",
                "mode": "accurate"
            }
    
    def convert_ocr_mode(self, pdf_path: str, output_path: str) -> Dict[str, Any]:
        """OCR conversion for scanned PDFs using pytesseract"""
        try:
            logger.info("Using OCR mode (pytesseract)")
            
            # Convert PDF pages to images
            images = convert_from_path(pdf_path, dpi=300)
            
            # Create Word document
            word_doc = Document()
            pages_processed = 0
            
            for page_num, image in enumerate(images):
                try:
                    # Add page break if not first page
                    if page_num > 0:
                        word_doc.add_page_break()
                    
                    # Perform OCR with better configuration
                    ocr_text = pytesseract.image_to_string(image, config='--psm 1 --oem 3 -c preserve_interword_spaces=1')
                    
                    if ocr_text.strip():
                        # Add OCR text with bullet point detection
                        self._add_text_with_bullets(word_doc, ocr_text.strip())
                    else:
                        # If no text found, add a note
                        paragraph = word_doc.add_paragraph()
                        paragraph.add_run(f"[Page {page_num + 1}: No text detected via OCR]")
                    
                    pages_processed += 1
                    
                except Exception as e:
                    logger.warning(f"Error processing page {page_num + 1} with OCR: {e}")
                    continue
            
            # Save document
            word_doc.save(output_path)
            
            return {
                "success": True,
                "message": "PDF converted using OCR mode (pytesseract)",
                "pages_processed": pages_processed,
                "mode": "ocr",
                "output_file": output_path
            }
            
        except Exception as e:
            logger.error(f"OCR mode conversion failed: {e}")
            return {
                "success": False,
                "error": f"OCR mode conversion failed: {str(e)}",
                "mode": "ocr"
            }
    
    def convert_hybrid_mode(self, pdf_path: str, output_path: str) -> Dict[str, Any]:
        """Hybrid conversion - images as backgrounds with overlaid text"""
        pdf_doc = None
        try:
            logger.info("Using Hybrid mode (pdf2image + PyMuPDF)")
            
            # Convert PDF pages to images
            page_images = convert_from_path(pdf_path, dpi=200)
            
            # Open PDF for text extraction
            pdf_doc = fitz.open(pdf_path)
            
            # Create Word document
            word_doc = Document()
            pages_processed = 0
            
            for page_num, page_image in enumerate(page_images):
                try:
                    # Add page break if not first page
                    if page_num > 0:
                        word_doc.add_page_break()
                    
                    # Save page image temporarily
                    img_path = os.path.join(self.temp_dir, f"page_{page_num}.png")
                    page_image.save(img_path, "PNG")
                    
                    # Add page image as background
                    paragraph = word_doc.add_paragraph()
                    run = paragraph.add_run()
                    
                    # Calculate image size (fit to page width)
                    img_width = Inches(6.5)  # Standard page width minus margins
                    img_height = Inches(img_width.inches * (page_image.height / page_image.width))
                    
                    run.add_picture(img_path, width=img_width, height=img_height)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Extract and overlay text if possible
                    if page_num < pdf_doc.page_count:
                        page = pdf_doc[page_num]
                        text = page.get_text()
                        if text.strip():
                            # Add text below image with bullet point detection
                            word_doc.add_paragraph()  # Add spacing
                            header_paragraph = word_doc.add_paragraph()
                            header_run = header_paragraph.add_run(f"[Extracted Text from Page {page_num + 1}]:")
                            header_run.font.size = Pt(8)
                            header_run.font.italic = True
                            self._add_text_with_bullets(word_doc, text.strip())
                    
                    pages_processed += 1
                    
                except Exception as e:
                    logger.warning(f"Error processing page {page_num + 1} in hybrid mode: {e}")
                    continue
            
            # Save document
            word_doc.save(output_path)
            if pdf_doc:
                pdf_doc.close()
            
            # Cleanup
            self._cleanup_temp_files()
            
            return {
                "success": True,
                "message": "PDF converted using Hybrid mode (images + text overlay)",
                "pages_processed": pages_processed,
                "mode": "hybrid",
                "output_file": output_path
            }
            
        except Exception as e:
            if pdf_doc:
                pdf_doc.close()
            logger.error(f"Hybrid mode conversion failed: {e}")
            return {
                "success": False,
                "error": f"Hybrid mode conversion failed: {str(e)}",
                "mode": "hybrid"
            }
    
    def convert_auto_mode(self, pdf_path: str, output_path: str) -> Dict[str, Any]:
        """Auto mode - intelligently chooses the best conversion method"""
        try:
            logger.info("Using Auto mode - detecting PDF type")
            
            pdf_type = self.detect_pdf_type(pdf_path)
            logger.info(f"Detected PDF type: {pdf_type}")
            
            if pdf_type == "scanned":
                # Use OCR for scanned PDFs
                logger.info("Auto mode: Using OCR for scanned PDF")
                result = self.convert_ocr_mode(pdf_path, output_path)
                result["auto_choice"] = "ocr"
                result["pdf_type"] = pdf_type
                return result
            
            elif pdf_type == "digital":
                # For digital PDFs, try accurate mode first, fallback to fast
                logger.info("Auto mode: Trying Accurate mode for digital PDF")
                result = self.convert_accurate_mode(pdf_path, output_path)
                
                if result["success"]:
                    result["auto_choice"] = "accurate"
                    result["pdf_type"] = pdf_type
                    return result
                else:
                    # Fallback to fast mode
                    logger.info("Auto mode: Falling back to Fast mode")
                    result = self.convert_fast_mode(pdf_path, output_path)
                    result["auto_choice"] = "fast_fallback"
                    result["pdf_type"] = pdf_type
                    return result
            
            else:  # mixed
                # For mixed content, use hybrid mode
                logger.info("Auto mode: Using Hybrid mode for mixed content PDF")
                result = self.convert_hybrid_mode(pdf_path, output_path)
                result["auto_choice"] = "hybrid"
                result["pdf_type"] = pdf_type
                return result
                
        except Exception as e:
            logger.error(f"Auto mode failed: {e}")
            return {
                "success": False,
                "error": f"Auto mode failed: {str(e)}",
                "mode": "auto"
            }
    
    def convert_pdf_to_word(self, pdf_path: str, output_path: str, mode: str = "auto") -> Dict[str, Any]:
        """Main conversion method with mode selection"""
        
        # Validate input
        if not os.path.exists(pdf_path):
            return {
                "success": False,
                "error": "Input PDF file not found",
                "mode": mode
            }
        
        # Convert mode to lowercase for consistency
        mode = mode.lower().strip()
        
        # Route to appropriate conversion method
        if mode == "fast":
            return self.convert_fast_mode(pdf_path, output_path)
        elif mode == "accurate":
            return self.convert_accurate_mode(pdf_path, output_path)
        elif mode == "hybrid":
            return self.convert_hybrid_mode(pdf_path, output_path)
        elif mode == "ocr":
            return self.convert_ocr_mode(pdf_path, output_path)
        elif mode == "auto":
            return self.convert_auto_mode(pdf_path, output_path)
        else:
            return {
                "success": False,
                "error": f"Unknown conversion mode: {mode}. Available modes: auto, fast, accurate, hybrid, ocr",
                "mode": mode
            }
    
    def _extract_formatted_text(self, page: fitz.Page, word_doc: Document):
        """Extract text with formatting for accurate mode with enhanced bullet point support"""
        try:
            text_dict = page.get_text("dict", flags=11)
            
            for block in text_dict["blocks"]:
                if "lines" in block:  # Text block
                    # Process lines to detect bullet points and structure
                    lines_text = []
                    for line in block["lines"]:
                        line_text = ""
                        line_spans = []
                        for span in line["spans"]:
                            if span["text"].strip():
                                line_text += span["text"]
                                line_spans.append(span)
                        if line_text.strip():
                            lines_text.append((line_text.strip(), line_spans))
                    
                    # Process each line with bullet point detection
                    for line_text, line_spans in lines_text:
                        is_bullet = self._is_bullet_point(line_text)
                        
                        if is_bullet:
                            # Create bullet point paragraph
                            paragraph = word_doc.add_paragraph(style='List Bullet')
                            # Remove bullet characters from text
                            clean_text = self._clean_bullet_text(line_text)
                        else:
                            # Regular paragraph
                            paragraph = word_doc.add_paragraph()
                            clean_text = line_text
                        
                        # Add text with original formatting
                        for span in line_spans:
                            if span["text"].strip():
                                # For bullet points, clean the text
                                span_text = self._clean_bullet_text(span["text"]) if is_bullet else span["text"]
                                if span_text.strip():
                                    run = paragraph.add_run(span_text)
                                    
                                    # Apply formatting
                                    font = run.font
                                    font.size = Pt(max(8, min(24, span["size"])))  # Reasonable font size limits
                                    font.name = span["font"] if span["font"] else "Calibri"
                                    
                                    # Handle formatting flags
                                    if span["flags"] & 2**4:  # Bold
                                        font.bold = True
                                    if span["flags"] & 2**1:  # Italic
                                        font.italic = True
                                        
                                    # Apply color if not black
                                    if span["color"] != 0:
                                        color_int = span["color"]
                                        r = (color_int >> 16) & 255
                                        g = (color_int >> 8) & 255
                                        b = color_int & 255
                                        font.color.rgb = RGBColor(r, g, b)
                        
        except Exception as e:
            logger.warning(f"Error extracting formatted text: {e}")
            # Fallback to basic text extraction with bullet point detection
            text = page.get_text()
            if text.strip():
                self._add_text_with_bullets(word_doc, text)
    
    def _extract_images_accurate(self, page: fitz.Page, word_doc: Document, pdf_doc):
        """Extract images for accurate mode"""
        try:
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    pix = fitz.Pixmap(pdf_doc, xref)
                    
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        # Convert to PIL Image
                        img_data = pix.tobytes("ppm")
                        pil_img = Image.open(io.BytesIO(img_data))
                        
                        # Save image temporarily
                        img_filename = f"image_{self.image_counter}.png"
                        img_path = os.path.join(self.temp_dir, img_filename)
                        pil_img.save(img_path, "PNG")
                        
                        # Add to Word document
                        paragraph = word_doc.add_paragraph()
                        run = paragraph.add_run()
                        
                        # Calculate appropriate size
                        width, height = pil_img.size
                        max_width = Inches(6)
                        
                        if width > max_width.emu:
                            ratio = height / width
                            width = max_width
                            height = Inches(max_width.inches * ratio)
                        else:
                            width = Inches(width / 96)
                            height = Inches(height / 96)
                        
                        run.add_picture(img_path, width=width, height=height)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        self.image_counter += 1
                        pil_img.close()
                    
                    pix = None  # Free memory
                    
                except Exception as e:
                    logger.warning(f"Error extracting image {img_index}: {e}")
                    continue
                    
        except Exception as e:
            logger.warning(f"Error in image extraction: {e}")
    
    def _extract_tables_accurate(self, page: fitz.Page, word_doc: Document):
        """Extract tables for accurate mode"""
        try:
            # Use PyMuPDF's table detection if available
            try:
                page_tables = page.find_tables()
                
                for table in page_tables:
                    table_data = table.extract()
                    if table_data:
                        # Create table in Word
                        rows = len(table_data)
                        cols = len(table_data[0]) if table_data else 0
                        
                        if rows > 0 and cols > 0:
                            word_table = word_doc.add_table(rows=rows, cols=cols)
                            word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            
                            # Fill table data
                            for row_idx, row_data in enumerate(table_data):
                                for col_idx, cell_data in enumerate(row_data):
                                    if col_idx < len(word_table.rows[row_idx].cells):
                                        cell = word_table.rows[row_idx].cells[col_idx]
                                        cell.text = str(cell_data) if cell_data else ""
                            
                            # Add spacing after table
                            word_doc.add_paragraph()
                            
            except AttributeError:
                # find_tables not available in this PyMuPDF version
                pass
                            
        except Exception as e:
            logger.warning(f"Error extracting tables: {e}")
    
    def _is_bullet_point(self, text: str) -> bool:
        """Detect if a line is a bullet point"""
        text = text.strip()
        if not text:
            return False
        
        # Common bullet point patterns
        bullet_patterns = [
            '•', '◦', '▪', '▫', '‣', '⁃',  # Unicode bullets
            '■', '□', '▲', '△', '●', '○',  # Geometric shapes
            '*', '-', '+',  # ASCII bullets
            '→', '➤', '►', '▶'  # Arrow bullets
        ]
        
        # Check if line starts with bullet character
        first_char = text[0]
        if first_char in bullet_patterns:
            return True
        
        # Check for numbered lists (1., 2., a., etc.)
        import re
        if re.match(r'^\d+[.):]\s+', text) or re.match(r'^[a-zA-Z][.):]\s+', text):
            return True
        
        # Check for indented text that might be a bullet
        if text.startswith('  ') or text.startswith('\t'):
            # Look for bullet-like patterns after whitespace
            stripped = text.lstrip()
            if stripped and stripped[0] in bullet_patterns:
                return True
        
        return False
    
    def _clean_bullet_text(self, text: str) -> str:
        """Remove bullet characters and clean bullet point text"""
        import re
        text = text.strip()
        
        # Remove common bullet characters from the beginning
        bullet_chars = ['•', '◦', '▪', '▫', '‣', '⁃', '■', '□', '▲', '△', '●', '○', '*', '-', '+', '→', '➤', '►', '▶']
        
        # Remove leading bullets
        if text and text[0] in bullet_chars:
            text = text[1:].strip()
        
        # Remove numbered list markers (1., 2., a., etc.)
        text = re.sub(r'^\d+[.):]\s*', '', text)
        text = re.sub(r'^[a-zA-Z][.):]\s*', '', text)
        
        # Remove extra whitespace
        text = text.strip()
        
        return text
    
    def _add_text_with_bullets(self, word_doc: Document, text: str):
        """Add text to document with bullet point detection"""
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if self._is_bullet_point(line):
                # Create bullet point
                paragraph = word_doc.add_paragraph(style='List Bullet')
                clean_text = self._clean_bullet_text(line)
                paragraph.add_run(clean_text)
            else:
                # Regular paragraph
                paragraph = word_doc.add_paragraph()
                paragraph.add_run(line)
    
    def _enhance_fast_mode_bullets(self, docx_path: str):
        """Post-process Fast mode output to enhance bullet point formatting"""
        try:
            # Open the document created by pdf2docx
            doc = Document(docx_path)
            
            # Process each paragraph to detect and fix bullet points
            paragraphs_to_modify = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text and self._is_bullet_point(text):
                    # Store paragraph info for modification
                    clean_text = self._clean_bullet_text(text)
                    paragraphs_to_modify.append((i, clean_text, paragraph))
            
            # Apply bullet formatting to detected bullet points
            for i, clean_text, paragraph in paragraphs_to_modify:
                # Clear existing content
                paragraph.clear()
                # Apply bullet style and add cleaned text
                paragraph.style = 'List Bullet'
                paragraph.add_run(clean_text)
            
            # Save the enhanced document
            doc.save(docx_path)
            
        except Exception as e:
            logger.warning(f"Error enhancing Fast mode bullets: {e}")
            # If enhancement fails, the original document remains
    
    def _cleanup_temp_files(self):
        """Clean up temporary files"""
        try:
            import shutil
            shutil.rmtree(self.temp_dir, ignore_errors=True)
        except Exception as e:
            logger.warning(f"Error cleaning up temp files: {e}")


def convert_pdf_to_word_enhanced(pdf_path: str, output_path: str, mode: str = "auto") -> Dict[str, Any]:
    """Enhanced main function to convert PDF to Word with mode selection"""
    converter = EnhancedPDFConverter()
    return converter.convert_pdf_to_word(pdf_path, output_path, mode)


# Backward compatibility
def convert_pdf_to_word_advanced(pdf_path: str, output_path: str) -> Dict[str, Any]:
    """Backward compatibility function - uses auto mode"""
    return convert_pdf_to_word_enhanced(pdf_path, output_path, "auto")


if __name__ == "__main__":
    # Test the converter
    import sys
    if len(sys.argv) > 2:
        pdf_file = sys.argv[1]
        word_file = sys.argv[2]
        mode = sys.argv[3] if len(sys.argv) > 3 else "auto"
        result = convert_pdf_to_word_enhanced(pdf_file, word_file, mode)
        print(result)