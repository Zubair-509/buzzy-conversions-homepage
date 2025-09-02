
#!/usr/bin/env python3

import os
import sys
import tempfile
import uuid
import argparse
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
import traceback
import time

# Core libraries
try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    print("Warning: docx2pdf not available. Will use fallback method.")

# Fallback libraries
try:
    from docx import Document
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    FALLBACK_AVAILABLE = True
except ImportError:
    FALLBACK_AVAILABLE = False
    print("Warning: Fallback libraries not available. Install python-docx and reportlab.")

def convert_with_docx2pdf(input_path: str, output_path: str) -> bool:
    """
    Convert DOCX to PDF using docx2pdf (primary method).
    This preserves full formatting, images, tables, and layout.
    """
    try:
        if not DOCX2PDF_AVAILABLE:
            return False
            
        print(f"Converting with docx2pdf: {os.path.basename(input_path)}")
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Convert using docx2pdf
        docx2pdf_convert(input_path, output_path)
        
        # Verify the output file was created
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            print(f"✓ Successfully converted using docx2pdf")
            return True
        else:
            print("✗ docx2pdf conversion failed - output file not created")
            return False
            
    except Exception as e:
        print(f"✗ docx2pdf conversion failed: {str(e)}")
        return False

def extract_docx_content(input_path: str) -> Dict[str, Any]:
    """
    Extract content from DOCX file for fallback conversion.
    """
    try:
        print(f"Extracting content from: {input_path}")
        doc = Document(input_path)
        content = {
            'paragraphs': [],
            'tables': [],
            'images': []
        }
        
        # Extract paragraphs with formatting
        for para in doc.paragraphs:
            # Include all paragraphs, even if they appear empty (might have formatting)
            para_text = para.text if para.text else ""
            
            para_info = {
                'text': para_text,
                'alignment': para.alignment,
                'runs': []
            }
            
            # Extract run-level formatting
            for run in para.runs:
                run_text = run.text if run.text else ""
                run_info = {
                    'text': run_text,
                    'bold': run.bold if run.bold is not None else False,
                    'italic': run.italic if run.italic is not None else False,
                    'underline': run.underline if run.underline is not None else False,
                    'font_name': run.font.name if run.font.name else 'Arial',
                    'font_size': run.font.size.pt if run.font.size else 12
                }
                para_info['runs'].append(run_info)
            
            # Add paragraph even if it seems empty (might contain formatting or be intentional)
            content['paragraphs'].append(para_info)
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip() if cell.text else ""
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            # Add table even if it appears empty
            content['tables'].append(table_data)
        
        # Debug: Print what we extracted
        print(f"Extracted {len(content['paragraphs'])} paragraphs and {len(content['tables'])} tables")
        
        # Check if we have any actual text content
        has_text = False
        for para in content['paragraphs']:
            if para['text'].strip():
                has_text = True
                break
        
        if not has_text:
            for table in content['tables']:
                for row in table:
                    for cell in row:
                        if cell.strip():
                            has_text = True
                            break
                    if has_text:
                        break
                if has_text:
                    break
        
        print(f"Document has text content: {has_text}")
        
        return content
        
    except Exception as e:
        print(f"Error extracting DOCX content: {e}")
        traceback.print_exc()
        return {'paragraphs': [], 'tables': [], 'images': []}

def convert_with_fallback(input_path: str, output_path: str) -> bool:
    """
    Convert DOCX to PDF using python-docx + reportlab (fallback method).
    This extracts content and creates a basic PDF with preserved formatting.
    """
    try:
        if not FALLBACK_AVAILABLE:
            print("✗ Fallback libraries not available")
            return False
            
        print(f"Converting with fallback method: {os.path.basename(input_path)}")
        
        # Extract content from DOCX
        content = extract_docx_content(input_path)
        
        # Check if we have any meaningful content
        has_content = False
        text_content = []
        
        # Collect all text content
        for para_info in content['paragraphs']:
            text = para_info['text'].strip()
            if text:
                has_content = True
                text_content.append(text)
        
        # Check tables for content
        for table_data in content['tables']:
            for row in table_data:
                for cell in row:
                    if cell.strip():
                        has_content = True
                        break
        
        # If no meaningful content found, create a simple document with a message
        if not has_content:
            print("⚠ No text content found, creating document with placeholder")
            text_content = ["This document appears to be empty or contains only formatting."]
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Create PDF
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Define custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.black
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=6,
            textColor=colors.black
        )
        
        # Add paragraphs
        if text_content:
            for text in text_content:
                # Escape HTML special characters for ReportLab
                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # Apply basic formatting based on content
                if len(text) < 100 and (text.isupper() or text.endswith(':')):
                    # Likely a heading
                    para = Paragraph(text, title_style)
                else:
                    # Regular paragraph
                    para = Paragraph(text, normal_style)
                
                story.append(para)
                story.append(Spacer(1, 6))
        else:
            # Add content from paragraphs directly
            for para_info in content['paragraphs']:
                text = para_info['text']
                if text.strip():
                    # Escape HTML special characters
                    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    para = Paragraph(text, normal_style)
                    story.append(para)
                    story.append(Spacer(1, 6))
        
        # Add tables
        for table_data in content['tables']:
            if table_data and any(any(cell.strip() for cell in row) for row in table_data):
                try:
                    # Create table with non-empty content
                    table = Table(table_data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    
                    story.append(table)
                    story.append(Spacer(1, 12))
                except Exception as table_error:
                    print(f"Error adding table: {table_error}")
                    # Add table content as paragraphs instead
                    for row in table_data:
                        row_text = " | ".join(cell.strip() for cell in row if cell.strip())
                        if row_text:
                            para = Paragraph(row_text, normal_style)
                            story.append(para)
        
        # If no content was added, add a default message
        if not story:
            para = Paragraph("Document converted successfully but appears to contain no text content.", normal_style)
            story.append(para)
        
        # Build PDF
        doc.build(story)
        
        # Verify the output file was created
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            print(f"✓ Successfully converted using fallback method")
            return True
        else:
            print("✗ Fallback conversion failed - output file not created")
            return False
            
    except Exception as e:
        print(f"✗ Fallback conversion failed: {str(e)}")
        traceback.print_exc()
        return False

def simple_text_extraction_fallback(input_path: str, output_path: str) -> bool:
    """
    Last resort: extract all text and create a simple PDF
    """
    try:
        print(f"Trying simple text extraction fallback for: {os.path.basename(input_path)}")
        
        doc = Document(input_path)
        all_text = []
        
        # Extract all text from paragraphs
        for para in doc.paragraphs:
            if para.text:
                all_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text:
                        row_text.append(cell.text)
                if row_text:
                    all_text.append(" | ".join(row_text))
        
        # If we still have no text, try to extract from runs
        if not all_text:
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.text:
                        all_text.append(run.text)
        
        if not all_text:
            all_text = ["This document appears to be empty or contains only images/formatting."]
        
        # Create simple PDF with extracted text
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        
        y_position = height - 50
        line_height = 14
        margin = 50
        
        c.setFont("Helvetica", 12)
        
        for text in all_text:
            if text.strip():
                # Split long lines
                words = text.split()
                line = ""
                
                for word in words:
                    test_line = line + " " + word if line else word
                    if c.stringWidth(test_line, "Helvetica", 12) < (width - 2 * margin):
                        line = test_line
                    else:
                        if line:
                            c.drawString(margin, y_position, line)
                            y_position -= line_height
                            line = word
                        else:
                            # Word is too long, break it
                            c.drawString(margin, y_position, word)
                            y_position -= line_height
                    
                    if y_position < margin:
                        c.showPage()
                        y_position = height - 50
                        c.setFont("Helvetica", 12)
                
                if line:
                    c.drawString(margin, y_position, line)
                    y_position -= line_height
                    
                y_position -= 5  # Extra space between paragraphs
                
                if y_position < margin:
                    c.showPage()
                    y_position = height - 50
                    c.setFont("Helvetica", 12)
        
        c.save()
        
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            print("✓ Simple text extraction successful")
            return True
        else:
            return False
            
    except Exception as e:
        print(f"Simple text extraction failed: {e}")
        return False

def convert_single_file(input_path: str, output_dir: str = None) -> Dict[str, Any]:
    """
    Convert a single DOCX file to PDF.
    """
    try:
        input_path = os.path.abspath(input_path)
        
        if not os.path.exists(input_path):
            return {
                'success': False,
                'error': f'Input file does not exist: {input_path}',
                'method': None
            }
        
        if not input_path.lower().endswith('.docx'):
            return {
                'success': False,
                'error': f'Invalid file type. Expected .docx, got: {os.path.splitext(input_path)[1]}',
                'method': None
            }
        
        # Determine output path
        if output_dir is None:
            output_dir = os.path.dirname(input_path)
        
        filename_without_ext = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.join(output_dir, f"{filename_without_ext}.pdf")
        
        # Try primary method first
        method_used = None
        success = False
        
        if DOCX2PDF_AVAILABLE:
            success = convert_with_docx2pdf(input_path, output_path)
            if success:
                method_used = "docx2pdf (primary)"
        
        # Try fallback method if primary failed
        if not success and FALLBACK_AVAILABLE:
            success = convert_with_fallback(input_path, output_path)
            if success:
                method_used = "python-docx + reportlab (fallback)"
        
        # Try simple text extraction as last resort
        if not success and FALLBACK_AVAILABLE:
            success = simple_text_extraction_fallback(input_path, output_path)
            if success:
                method_used = "simple text extraction (last resort)"
        
        if not success:
            return {
                'success': False,
                'error': 'All conversion methods failed. The DOCX file may be corrupted or have an unsupported format.',
                'method': None
            }
        
        file_size = os.path.getsize(output_path)
        
        return {
            'success': True,
            'input_file': os.path.basename(input_path),
            'output_file': os.path.basename(output_path),
            'output_path': output_path,
            'method': method_used,
            'file_size': file_size
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': f'Conversion error: {str(e)}',
            'method': None
        }

def convert_batch(input_dir: str, output_dir: str = None) -> Dict[str, Any]:
    """
    Convert multiple DOCX files in a directory to PDF.
    """
    try:
        input_dir = os.path.abspath(input_dir)
        
        if not os.path.exists(input_dir):
            return {
                'success': False,
                'error': f'Input directory does not exist: {input_dir}',
                'results': []
            }
        
        if output_dir is None:
            output_dir = os.path.join(input_dir, 'pdf_output')
        
        os.makedirs(output_dir, exist_ok=True)
        
        # Find all DOCX files
        docx_files = []
        for file in os.listdir(input_dir):
            if file.lower().endswith('.docx') and not file.startswith('~'):  # Skip temp files
                docx_files.append(os.path.join(input_dir, file))
        
        if not docx_files:
            return {
                'success': False,
                'error': f'No DOCX files found in: {input_dir}',
                'results': []
            }
        
        print(f"Found {len(docx_files)} DOCX files to convert")
        print(f"Output directory: {output_dir}")
        
        results = []
        successful_conversions = 0
        
        for i, docx_file in enumerate(docx_files, 1):
            print(f"\n[{i}/{len(docx_files)}] Processing: {os.path.basename(docx_file)}")
            
            result = convert_single_file(docx_file, output_dir)
            results.append(result)
            
            if result['success']:
                successful_conversions += 1
                print(f"    ✓ Converted using {result['method']}")
            else:
                print(f"    ✗ Failed: {result['error']}")
        
        return {
            'success': True,
            'message': f'Batch conversion completed. {successful_conversions}/{len(docx_files)} files converted successfully.',
            'total_files': len(docx_files),
            'successful_conversions': successful_conversions,
            'results': results
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': f'Batch conversion error: {str(e)}',
            'results': []
        }

def check_dependencies():
    """
    Check which conversion methods are available.
    """
    print("🔍 Checking dependencies...")
    
    if DOCX2PDF_AVAILABLE:
        print("✓ docx2pdf: Available (Primary method)")
    else:
        print("✗ docx2pdf: Not available")
        print("  Install with: pip install docx2pdf")
    
    if FALLBACK_AVAILABLE:
        print("✓ python-docx + reportlab: Available (Fallback method)")
    else:
        print("✗ python-docx + reportlab: Not available")
        print("  Install with: pip install python-docx reportlab")
    
    if not DOCX2PDF_AVAILABLE and not FALLBACK_AVAILABLE:
        print("\n❌ No conversion methods available. Please install dependencies.")
        return False
    
    print("\n✅ Ready for conversion!")
    return True

def main():
    """
    Main function with CLI support.
    """
    parser = argparse.ArgumentParser(
        description="Convert Word (.docx) files to PDF with high accuracy",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python word_to_pdf_converter.py document.docx
  python word_to_pdf_converter.py documents_folder/
  python word_to_pdf_converter.py document.docx --output ./pdfs/
  python word_to_pdf_converter.py --check-deps
        """
    )
    
    parser.add_argument(
        'input',
        nargs='?',
        help='Input DOCX file or directory containing DOCX files'
    )
    
    parser.add_argument(
        '--output', '-o',
        help='Output directory for PDF files'
    )
    
    parser.add_argument(
        '--check-deps',
        action='store_true',
        help='Check available dependencies and exit'
    )
    
    args = parser.parse_args()
    
    # Check dependencies
    if args.check_deps:
        check_dependencies()
        return
    
    if not args.input:
        parser.print_help()
        return
    
    # Check if conversion methods are available
    if not check_dependencies():
        sys.exit(1)
    
    input_path = args.input
    
    # Determine if input is file or directory
    if os.path.isfile(input_path):
        # Single file conversion
        print(f"\n🔄 Converting single file: {os.path.basename(input_path)}")
        result = convert_single_file(input_path, args.output)
        
        if result['success']:
            print(f"\n✅ Conversion successful!")
            print(f"   Method: {result['method']}")
            print(f"   Output: {result['output_file']}")
            print(f"   Size: {result['file_size']} bytes")
        else:
            print(f"\n❌ Conversion failed: {result['error']}")
            sys.exit(1)
            
    elif os.path.isdir(input_path):
        # Batch conversion
        print(f"\n🔄 Starting batch conversion from: {input_path}")
        result = convert_batch(input_path, args.output)
        
        if result['success']:
            print(f"\n✅ {result['message']}")
            print(f"   Success rate: {result['successful_conversions']}/{result['total_files']}")
        else:
            print(f"\n❌ Batch conversion failed: {result['error']}")
            sys.exit(1)
    else:
        print(f"❌ Input path does not exist: {input_path}")
        sys.exit(1)

if __name__ == '__main__':
    main()
