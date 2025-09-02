
#!/usr/bin/env python3

import os
import tempfile
import uuid
import io
import base64
from pathlib import Path
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from pdf2docx import Converter
from werkzeug.utils import secure_filename
import traceback
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import pytesseract
from PIL import Image
import threading
import time
from typing import List, Dict, Any, Tuple, Optional

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
ALLOWED_EXTENSIONS = {'pdf'}

# Get the directory where this script is located
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)

# Use absolute paths relative to project root
UPLOAD_FOLDER = os.path.join(PROJECT_ROOT, 'uploads')
OUTPUT_FOLDER = os.path.join(PROJECT_ROOT, 'outputs')

# Create directories if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def cleanup_file(filepath):
    """Safely remove a file if it exists"""
    try:
        if os.path.exists(filepath):
            os.remove(filepath)
    except Exception as e:
        print(f"Warning: Could not remove file {filepath}: {e}")

def is_scanned_pdf(pdf_path: str) -> bool:
    """Detect if PDF is scanned (image-based) or digital (text-based)"""
    try:
        doc = fitz.open(pdf_path)
        text_ratio = 0
        total_pages = len(doc)
        
        for page_num in range(min(3, total_pages)):  # Check first 3 pages
            page = doc[page_num]
            text = page.get_text().strip()
            images = page.get_images()
            
            if len(text) > 100:  # Has substantial text
                text_ratio += 1
            elif len(images) > 0:  # Has images but little text
                continue
                
        doc.close()
        
        # If less than 50% of checked pages have text, consider it scanned
        return (text_ratio / min(3, total_pages)) < 0.5
        
    except Exception as e:
        print(f"Error detecting PDF type: {e}")
        return False

def extract_text_with_formatting(pdf_path: str) -> List[Dict[str, Any]]:
    """Extract text with formatting information using PyMuPDF"""
    doc = fitz.open(pdf_path)
    pages_data = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Get text blocks with formatting
        blocks = page.get_text("dict")
        
        page_data = {
            'page_num': page_num,
            'blocks': [],
            'images': [],
            'tables': []
        }
        
        # Extract text blocks
        for block in blocks["blocks"]:
            if "lines" in block:
                for line in block["lines"]:
                    for span in line["spans"]:
                        text_info = {
                            'text': span["text"],
                            'bbox': span["bbox"],
                            'font': span["font"],
                            'size': span["size"],
                            'flags': span["flags"],  # Bold, italic flags
                            'color': span["color"]
                        }
                        page_data['blocks'].append(text_info)
        
        # Extract images
        image_list = page.get_images()
        for img_index, img in enumerate(image_list):
            try:
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    img_data = pix.tobytes("png")
                    page_data['images'].append({
                        'index': img_index,
                        'data': base64.b64encode(img_data).decode(),
                        'bbox': page.get_image_bbox(img)
                    })
                pix = None
            except Exception as e:
                print(f"Error extracting image: {e}")
        
        # Detect tables (simplified table detection)
        tables = detect_tables(page)
        page_data['tables'] = tables
        
        pages_data.append(page_data)
    
    doc.close()
    return pages_data

def detect_tables(page) -> List[Dict[str, Any]]:
    """Simple table detection based on text alignment"""
    blocks = page.get_text("dict")
    tables = []
    
    # Group text by vertical position to find potential table rows
    text_lines = []
    for block in blocks["blocks"]:
        if "lines" in block:
            for line in block["lines"]:
                line_text = ""
                bbox = None
                for span in line["spans"]:
                    line_text += span["text"]
                    if bbox is None:
                        bbox = list(span["bbox"])
                    else:
                        bbox[2] = max(bbox[2], span["bbox"][2])
                        bbox[3] = max(bbox[3], span["bbox"][3])
                
                if line_text.strip():
                    text_lines.append({
                        'text': line_text.strip(),
                        'bbox': bbox,
                        'y': bbox[1]
                    })
    
    # Sort by vertical position
    text_lines.sort(key=lambda x: x['y'])
    
    # Simple table detection: look for aligned text patterns
    potential_rows = []
    current_row = []
    last_y = None
    
    for line in text_lines:
        if last_y is None or abs(line['y'] - last_y) < 5:  # Same row
            current_row.append(line)
        else:
            if len(current_row) > 1:  # Potential table row
                potential_rows.append(current_row)
            current_row = [line]
        last_y = line['y']
    
    if len(current_row) > 1:
        potential_rows.append(current_row)
    
    # If we have multiple rows with similar column structure, it's likely a table
    if len(potential_rows) > 2:
        tables.append({
            'rows': potential_rows,
            'bbox': [
                min(row[0]['bbox'][0] for row in potential_rows),
                min(row[0]['bbox'][1] for row in potential_rows),
                max(row[-1]['bbox'][2] for row in potential_rows),
                max(row[-1]['bbox'][3] for row in potential_rows)
            ]
        })
    
    return tables

def ocr_extract_text(pdf_path: str) -> str:
    """Extract text from scanned PDF using OCR"""
    try:
        doc = fitz.open(pdf_path)
        all_text = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Convert page to image
            mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better OCR
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            
            # Convert to PIL Image for OCR
            image = Image.open(io.BytesIO(img_data))
            
            # Perform OCR
            text = pytesseract.image_to_string(image, lang='eng')
            all_text.append(f"--- Page {page_num + 1} ---\n{text}\n")
            
        doc.close()
        return "\n".join(all_text)
        
    except Exception as e:
        print(f"OCR extraction failed: {e}")
        return ""

def build_word_document(pages_data: List[Dict[str, Any]], output_path: str, mode: str = 'standard'):
    """Build Word document from extracted data"""
    doc = Document()
    
    for page_data in pages_data:
        # Add page break between pages (except first)
        if page_data['page_num'] > 0:
            doc.add_page_break()
        
        # Sort blocks by vertical position for reading order
        blocks = sorted(page_data['blocks'], key=lambda x: (x['bbox'][1], x['bbox'][0]))
        
        # Group blocks into paragraphs
        paragraphs = group_blocks_into_paragraphs(blocks)
        
        # Add paragraphs to document
        for para_blocks in paragraphs:
            if not para_blocks:
                continue
                
            paragraph = doc.add_paragraph()
            
            for block in para_blocks:
                run = paragraph.add_run(block['text'])
                
                # Apply formatting
                apply_text_formatting(run, block)
        
        # Add tables
        for table_data in page_data['tables']:
            add_table_to_document(doc, table_data)
        
        # Add images
        for img_data in page_data['images']:
            add_image_to_document(doc, img_data)
    
    doc.save(output_path)

def group_blocks_into_paragraphs(blocks: List[Dict[str, Any]]) -> List[List[Dict[str, Any]]]:
    """Group text blocks into logical paragraphs"""
    if not blocks:
        return []
    
    paragraphs = []
    current_paragraph = []
    last_bottom = None
    
    for block in blocks:
        bbox = block['bbox']
        top = bbox[1]
        bottom = bbox[3]
        
        # Start new paragraph if there's significant vertical gap
        if last_bottom is not None and top - last_bottom > 10:
            if current_paragraph:
                paragraphs.append(current_paragraph)
                current_paragraph = []
        
        current_paragraph.append(block)
        last_bottom = bottom
    
    if current_paragraph:
        paragraphs.append(current_paragraph)
    
    return paragraphs

def apply_text_formatting(run, block: Dict[str, Any]):
    """Apply formatting to text run based on extracted information"""
    flags = block.get('flags', 0)
    size = block.get('size', 12)
    
    # Apply font size
    run.font.size = Pt(size)
    
    # Apply bold
    if flags & 2**4:  # Bold flag
        run.font.bold = True
    
    # Apply italic
    if flags & 2**1:  # Italic flag
        run.font.italic = True
    
    # Apply font family if available
    font_name = block.get('font', '').split('+')[-1]  # Remove subset prefix
    if font_name:
        run.font.name = font_name

def add_table_to_document(doc: Document, table_data: Dict[str, Any]):
    """Add table to Word document"""
    rows = table_data.get('rows', [])
    if not rows:
        return
    
    # Create table with detected dimensions
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < max_cols:
                cell = table.cell(row_idx, col_idx)
                cell.text = cell_data.get('text', '')

def add_image_to_document(doc: Document, img_data: Dict[str, Any]):
    """Add image to Word document"""
    try:
        # Decode base64 image data
        img_bytes = base64.b64decode(img_data['data'])
        
        # Create image stream
        img_stream = io.BytesIO(img_bytes)
        
        # Add image to document
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        
        # Calculate appropriate size (max 6 inches width)
        max_width = Inches(6)
        run.add_picture(img_stream, width=max_width)
        
    except Exception as e:
        print(f"Error adding image: {e}")

def convert_with_pdf2docx(pdf_path: str, docx_path: str) -> bool:
    """Convert using pdf2docx library (standard mode)"""
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0)
        cv.close()
        return True
    except Exception as e:
        print(f"pdf2docx conversion failed: {e}")
        return False

def convert_with_advanced_method(pdf_path: str, docx_path: str, mode: str = 'accurate') -> bool:
    """Convert using advanced PyMuPDF + python-docx method"""
    try:
        # Extract structured data
        pages_data = extract_text_with_formatting(pdf_path)
        
        if not pages_data:
            return False
        
        # Build Word document
        build_word_document(pages_data, docx_path, mode)
        return True
        
    except Exception as e:
        print(f"Advanced conversion failed: {e}")
        return False

def convert_with_ocr(pdf_path: str, docx_path: str) -> bool:
    """Convert scanned PDF using OCR"""
    try:
        # Extract text using OCR
        text = ocr_extract_text(pdf_path)
        
        if not text.strip():
            return False
        
        # Create simple Word document with OCR text
        doc = Document()
        
        # Split by page markers
        pages = text.split('--- Page')
        for page_text in pages:
            if page_text.strip():
                # Add page content
                paragraphs = page_text.strip().split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
                
                # Add page break (except for last page)
                if page_text != pages[-1]:
                    doc.add_page_break()
        
        doc.save(docx_path)
        return True
        
    except Exception as e:
        print(f"OCR conversion failed: {e}")
        return False

def create_hybrid_document(pdf_path: str, docx_path: str) -> bool:
    """Create hybrid document with background images and overlay text"""
    try:
        doc = Document()
        pdf_doc = fitz.open(pdf_path)
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            
            # Convert page to high-resolution image
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            
            # Add page image as background
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            img_stream = io.BytesIO(img_data)
            run.add_picture(img_stream, width=Inches(7.5))
            
            # Extract and overlay text
            text_blocks = page.get_text("dict")
            for block in text_blocks["blocks"]:
                if "lines" in block:
                    block_text = ""
                    for line in block["lines"]:
                        for span in line["spans"]:
                            block_text += span["text"]
                    
                    if block_text.strip():
                        # Add text as transparent overlay (simplified)
                        text_para = doc.add_paragraph(block_text.strip())
                        text_para.paragraph_format.space_after = Pt(6)
            
            # Add page break
            if page_num < len(pdf_doc) - 1:
                doc.add_page_break()
        
        pdf_doc.close()
        doc.save(docx_path)
        return True
        
    except Exception as e:
        print(f"Hybrid conversion failed: {e}")
        return False

def batch_convert_pdfs(pdf_paths: List[str], output_dir: str, mode: str = 'standard') -> Dict[str, Any]:
    """Convert multiple PDFs in batch"""
    results = []
    
    for pdf_path in pdf_paths:
        try:
            filename = os.path.basename(pdf_path)
            name_without_ext = os.path.splitext(filename)[0]
            unique_id = str(uuid.uuid4())
            
            docx_filename = f"{name_without_ext}_{unique_id}.docx"
            docx_path = os.path.join(output_dir, docx_filename)
            
            # Convert based on mode
            success = False
            if mode == 'fast':
                success = convert_with_pdf2docx(pdf_path, docx_path)
            elif mode == 'accurate':
                success = convert_with_advanced_method(pdf_path, docx_path, mode)
            elif mode == 'hybrid':
                success = create_hybrid_document(pdf_path, docx_path)
            
            results.append({
                'original_file': filename,
                'output_file': docx_filename,
                'download_id': unique_id,
                'success': success
            })
            
        except Exception as e:
            results.append({
                'original_file': os.path.basename(pdf_path),
                'success': False,
                'error': str(e)
            })
    
    return {'results': results}

@app.route('/api/convert/pdf-to-word', methods=['POST'])
def convert_pdf_to_word():
    try:
        # Check if file is present in request
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        conversion_mode = request.form.get('mode', 'standard')  # standard, fast, accurate, hybrid
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type. Only PDF files are allowed.'}), 400

        # Generate unique filename to avoid conflicts
        unique_id = str(uuid.uuid4())
        original_filename = secure_filename(file.filename or 'uploaded_file.pdf')
        filename_without_ext = os.path.splitext(original_filename)[0]

        # Save uploaded file
        pdf_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}_{original_filename}")
        file.save(pdf_path)

        # Generate output path
        docx_filename = f"{filename_without_ext}_{unique_id}.docx"
        docx_path = os.path.join(OUTPUT_FOLDER, docx_filename)

        # Convert PDF to Word based on mode and PDF type
        try:
            success = False
            conversion_method = "unknown"
            
            if conversion_mode == 'fast':
                # Fast mode: use pdf2docx
                success = convert_with_pdf2docx(pdf_path, docx_path)
                conversion_method = "pdf2docx (fast)"
                
            elif conversion_mode == 'hybrid':
                # Hybrid mode: background image + text overlay
                success = create_hybrid_document(pdf_path, docx_path)
                conversion_method = "hybrid (image + text)"
                
            else:
                # Auto-detect PDF type and choose best method
                is_scanned = is_scanned_pdf(pdf_path)
                
                if is_scanned:
                    # Use OCR for scanned PDFs
                    success = convert_with_ocr(pdf_path, docx_path)
                    conversion_method = "OCR (scanned PDF)"
                    
                    # Fallback to advanced method if OCR fails
                    if not success:
                        success = convert_with_advanced_method(pdf_path, docx_path, conversion_mode)
                        conversion_method = "PyMuPDF + python-docx (OCR fallback)"
                else:
                    # Try pdf2docx first for digital PDFs
                    success = convert_with_pdf2docx(pdf_path, docx_path)
                    conversion_method = "pdf2docx (digital)"
                    
                    # Fallback to advanced method if pdf2docx fails
                    if not success or conversion_mode == 'accurate':
                        success = convert_with_advanced_method(pdf_path, docx_path, conversion_mode)
                        conversion_method = "PyMuPDF + python-docx (advanced)"

            # Cleanup input file
            cleanup_file(pdf_path)

            # Check if conversion was successful
            if not success or not os.path.exists(docx_path):
                return jsonify({'error': 'Conversion failed. Output file was not created.'}), 500

            return jsonify({
                'success': True,
                'message': f'PDF converted to Word successfully using {conversion_method}',
                'download_id': unique_id,
                'filename': docx_filename,
                'conversion_method': conversion_method,
                'file_size': os.path.getsize(docx_path)
            })

        except Exception as conv_error:
            # Cleanup files on conversion error
            cleanup_file(pdf_path)
            cleanup_file(docx_path)
            print(f"Conversion error: {conv_error}")
            print(traceback.format_exc())
            return jsonify({'error': f'Conversion failed: {str(conv_error)}'}), 500

    except Exception as e:
        print(f"General error: {e}")
        print(traceback.format_exc())
        return jsonify({'error': f'Server error: {str(e)}'}), 500

@app.route('/api/convert/pdf-to-word/batch', methods=['POST'])
def batch_convert_pdf_to_word():
    """Handle batch conversion of multiple PDFs"""
    try:
        files = request.files.getlist('files')
        conversion_mode = request.form.get('mode', 'standard')
        
        if not files:
            return jsonify({'error': 'No files provided'}), 400
        
        pdf_paths = []
        file_info = []
        
        # Save all uploaded files
        for file in files:
            if file.filename and allowed_file(file.filename):
                unique_id = str(uuid.uuid4())
                original_filename = secure_filename(file.filename)
                pdf_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}_{original_filename}")
                file.save(pdf_path)
                
                pdf_paths.append(pdf_path)
                file_info.append({
                    'original_name': original_filename,
                    'unique_id': unique_id,
                    'path': pdf_path
                })
        
        if not pdf_paths:
            return jsonify({'error': 'No valid PDF files found'}), 400
        
        # Process batch conversion
        results = batch_convert_pdfs(pdf_paths, OUTPUT_FOLDER, conversion_mode)
        
        # Cleanup input files
        for pdf_path in pdf_paths:
            cleanup_file(pdf_path)
        
        return jsonify({
            'success': True,
            'message': f'Batch conversion completed. {len(results["results"])} files processed.',
            'results': results['results']
        })
        
    except Exception as e:
        print(f"Batch conversion error: {e}")
        print(traceback.format_exc())
        return jsonify({'error': f'Batch conversion failed: {str(e)}'}), 500

@app.route('/api/download/<download_id>/<filename>', methods=['GET'])
def download_file(download_id, filename):
    try:
        file_path = os.path.join(OUTPUT_FOLDER, filename)

        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found'}), 404

        # Create response with the file
        response = send_file(
            file_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        # Schedule file cleanup after download
        def cleanup_after_download():
            def delayed_cleanup():
                time.sleep(30)  # Wait 30 seconds before cleanup
                cleanup_file(file_path)
            threading.Thread(target=delayed_cleanup).start()

        cleanup_after_download()

        return response

    except Exception as e:
        print(f"Download error: {e}")
        print(traceback.format_exc())
        return jsonify({'error': f'Download failed: {str(e)}'}), 500

# Import Word to PDF converter
try:
    from word_to_pdf_converter import convert_single_file as word_convert_single, convert_batch as word_convert_batch
    from word_to_pdf_converter import DOCX2PDF_AVAILABLE, FALLBACK_AVAILABLE
    WORD_TO_PDF_AVAILABLE = True
except ImportError:
    WORD_TO_PDF_AVAILABLE = False
    print("Warning: Word to PDF converter not available")

@app.route('/api/convert/word-to-pdf', methods=['POST'])
def convert_word_to_pdf_endpoint():
    """Convert Word documents to PDF"""
    if not WORD_TO_PDF_AVAILABLE:
        return jsonify({'error': 'Word to PDF conversion not available. Missing dependencies.'}), 503
    
    try:
        # Check if file is present in request
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not file.filename.lower().endswith('.docx'):
            return jsonify({'error': 'Invalid file type. Only DOCX files are allowed.'}), 400

        # Generate unique filename to avoid conflicts
        unique_id = str(uuid.uuid4())
        original_filename = secure_filename(file.filename or 'uploaded_file.docx')

        # Save uploaded file
        docx_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}_{original_filename}")
        file.save(docx_path)

        try:
            # Convert DOCX to PDF
            result = word_convert_single(docx_path, OUTPUT_FOLDER)
            
            # Cleanup input file
            cleanup_file(docx_path)

            if not result['success']:
                return jsonify({'error': f"Conversion failed: {result['error']}"}), 500

            return jsonify({
                'success': True,
                'message': f'Word document converted to PDF successfully using {result["method"]}',
                'download_id': unique_id,
                'filename': result['output_file'],
                'conversion_method': result['method'],
                'file_size': result['file_size']
            })

        except Exception as conv_error:
            # Cleanup files on conversion error
            cleanup_file(docx_path)
            print(f"Word to PDF conversion error: {conv_error}")
            print(traceback.format_exc())
            return jsonify({'error': f'Conversion failed: {str(conv_error)}'}), 500

    except Exception as e:
        print(f"Word to PDF general error: {e}")
        print(traceback.format_exc())
        return jsonify({'error': f'Server error: {str(e)}'}), 500

@app.route('/api/convert/word-to-pdf/batch', methods=['POST'])
def batch_convert_word_to_pdf_endpoint():
    """Handle batch conversion of multiple DOCX files"""
    if not WORD_TO_PDF_AVAILABLE:
        return jsonify({'error': 'Word to PDF conversion not available. Missing dependencies.'}), 503
    
    try:
        files = request.files.getlist('files')
        
        if not files:
            return jsonify({'error': 'No files provided'}), 400
        
        docx_paths = []
        file_info = []
        
        # Save all uploaded files
        for file in files:
            if file.filename and file.filename.lower().endswith('.docx'):
                unique_id = str(uuid.uuid4())
                original_filename = secure_filename(file.filename)
                docx_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}_{original_filename}")
                file.save(docx_path)
                
                docx_paths.append(docx_path)
                file_info.append({
                    'original_name': original_filename,
                    'unique_id': unique_id,
                    'path': docx_path
                })
        
        if not docx_paths:
            return jsonify({'error': 'No valid DOCX files found'}), 400
        
        # Process batch conversion
        results = []
        successful_conversions = 0
        
        for info in file_info:
            try:
                result = word_convert_single(info['path'], OUTPUT_FOLDER)
                
                if result['success']:
                    successful_conversions += 1
                    results.append({
                        'original_file': info['original_name'],
                        'output_file': result['output_file'],
                        'download_id': info['unique_id'],
                        'success': True,
                        'method': result['method'],
                        'file_size': result['file_size']
                    })
                else:
                    results.append({
                        'original_file': info['original_name'],
                        'success': False,
                        'error': result['error']
                    })
                    
            except Exception as e:
                results.append({
                    'original_file': info['original_name'],
                    'success': False,
                    'error': str(e)
                })
        
        # Cleanup input files
        for docx_path in docx_paths:
            cleanup_file(docx_path)
        
        return jsonify({
            'success': True,
            'message': f'Batch conversion completed. {successful_conversions}/{len(results)} files converted successfully.',
            'total_files': len(results),
            'successful_conversions': successful_conversions,
            'results': results
        })
        
    except Exception as e:
        print(f"Word to PDF batch conversion error: {e}")
        print(traceback.format_exc())
        return jsonify({'error': f'Batch conversion failed: {str(e)}'}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    # Check Word to PDF availability
    word_pdf_methods = []
    if WORD_TO_PDF_AVAILABLE:
        if DOCX2PDF_AVAILABLE:
            word_pdf_methods.append('docx2pdf (primary)')
        if FALLBACK_AVAILABLE:
            word_pdf_methods.append('python-docx + reportlab (fallback)')
    
    return jsonify({
        'status': 'healthy', 
        'service': 'Enhanced PDF & Word Converter',
        'features': [
            'Advanced formatting preservation',
            'OCR for scanned PDFs', 
            'Multi-column layout detection',
            'Table and image extraction',
            'Batch processing',
            'Hybrid mode with background images',
            'Word to PDF conversion',
            'High-accuracy DOCX processing'
        ],
        'word_to_pdf_methods': word_pdf_methods if WORD_TO_PDF_AVAILABLE else ['Not available']
    })

@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File too large. Maximum size is 16MB.'}), 413

if __name__ == '__main__':
    # Set maximum file size to 16MB
    app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

    # Run the Flask app
    port = int(os.environ.get('PORT', os.environ.get('PYTHON_PORT', 8001)))
    app.run(host='0.0.0.0', port=port, debug=False)
