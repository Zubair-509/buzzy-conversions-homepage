
#!/usr/bin/env python3
"""
Enhanced Word to PDF Converter
Converts Word documents to PDF while preserving formatting, layout, images, and tables.
Uses multiple methods for optimal conversion quality with advanced formatting preservation.
"""

import os
import tempfile
import traceback
import shutil
from pathlib import Path
from typing import Dict, Any, Optional, List
import uuid
import subprocess
import platform

try:
    from docx2pdf import convert as docx2pdf_convert
    print("docx2pdf successfully imported")
except ImportError as e:
    print(f"Warning: docx2pdf not available - {e}")
    docx2pdf_convert = None

try:
    import pypandoc
    print("pypandoc successfully imported")
except ImportError as e:
    print(f"Warning: pypandoc not available - {e}")
    pypandoc = None

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    print("python-docx successfully imported")
except ImportError as e:
    print(f"Warning: python-docx not available - {e}")
    Document = None

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    print("reportlab successfully imported")
except ImportError as e:
    print(f"Warning: reportlab not available - {e}")
    canvas = None

try:
    from PIL import Image
    print("Pillow successfully imported")
except ImportError as e:
    print(f"Warning: Pillow not available - {e}")
    Image = None

class WordToPDFConverter:
    """Enhanced Word to PDF converter with advanced formatting preservation"""
    
    def __init__(self, temp_dir: str = None):
        self.temp_dir = temp_dir or tempfile.gettempdir()
        self.conversion_cache = {}
        
    def convert_word_to_pdf(self, word_path: str, output_path: str = None) -> Dict[str, Any]:
        """
        Convert Word document to PDF with enhanced formatting preservation
        
        Args:
            word_path: Path to input Word file
            output_path: Path for output PDF file (optional)
            
        Returns:
            Dictionary with conversion results
        """
        try:
            if not os.path.exists(word_path):
                return {"success": False, "error": f"Word file not found: {word_path}"}
            
            # Validate file extension
            if not word_path.lower().endswith(('.docx', '.doc')):
                return {"success": False, "error": "File must be a Word document (.doc or .docx)"}
            
            # Generate output path if not provided
            if not output_path:
                base_name = os.path.splitext(os.path.basename(word_path))[0]
                output_path = os.path.join(self.temp_dir, f"{base_name}_converted.pdf")
            
            print(f"Starting Word to PDF conversion: {word_path} -> {output_path}")
            
            # Get document metadata
            metadata = self._get_word_metadata(word_path)
            
            # Method 1: docx2pdf (Windows/Mac optimized)
            if docx2pdf_convert and self._convert_with_docx2pdf(word_path, output_path):
                print("docx2pdf conversion method succeeded")
                # Safely get file size
                file_size = 0
                try:
                    if os.path.exists(output_path):
                        file_size = os.path.getsize(output_path)
                except Exception as e:
                    print(f"Warning: Could not get file size: {e}")
                
                return {
                    "success": True,
                    "output_path": output_path,
                    "filename": os.path.basename(output_path),
                    "method": "docx2pdf_native",
                    "metadata": metadata,
                    "file_size": file_size
                }
            
            # Method 2: LibreOffice command line
            if self._convert_with_libreoffice(word_path, output_path):
                print("LibreOffice conversion method succeeded")
                # Safely get file size
                file_size = 0
                try:
                    if os.path.exists(output_path):
                        file_size = os.path.getsize(output_path)
                except Exception as e:
                    print(f"Warning: Could not get file size: {e}")
                
                return {
                    "success": True,
                    "output_path": output_path,
                    "filename": os.path.basename(output_path),
                    "method": "libreoffice_headless",
                    "metadata": metadata,
                    "file_size": file_size
                }
            
            # Method 3: Pandoc conversion
            if pypandoc and self._convert_with_pandoc(word_path, output_path):
                print("Pandoc conversion method succeeded")
                return {
                    "success": True,
                    "output_path": output_path,
                    "filename": os.path.basename(output_path),
                    "method": "pandoc_enhanced",
                    "metadata": metadata,
                    "file_size": os.path.getsize(output_path)
                }
            
            # Method 4: Manual conversion with ReportLab (fallback)
            if self._convert_with_reportlab_fallback(word_path, output_path):
                print("ReportLab fallback conversion method succeeded")
                return {
                    "success": True,
                    "output_path": output_path,
                    "filename": os.path.basename(output_path),
                    "method": "reportlab_manual",
                    "metadata": metadata,
                    "file_size": os.path.getsize(output_path)
                }
            
            return {"success": False, "error": "All conversion methods failed. Please ensure LibreOffice is installed or try a different document."}
            
        except Exception as e:
            error_msg = f"Word to PDF conversion failed: {str(e)}"
            print(f"Error: {error_msg}")
            print(traceback.format_exc())
            return {"success": False, "error": error_msg}
    
    def _get_word_metadata(self, word_path: str) -> Dict[str, Any]:
        """Extract metadata from Word document"""
        metadata = {
            "pages": 0,
            "has_images": False,
            "has_tables": False,
            "title": "",
            "author": "",
            "subject": "",
            "creator": "Word to PDF Converter"
        }
        
        try:
            if Document and word_path.lower().endswith('.docx'):
                doc = Document(word_path)
                
                # Count paragraphs and estimate pages
                paragraphs = len(doc.paragraphs)
                metadata["pages"] = max(1, paragraphs // 30)  # Rough estimate
                
                # Check for tables
                if doc.tables:
                    metadata["has_tables"] = True
                
                # Check for images
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        metadata["has_images"] = True
                        break
                
                # Get core properties
                if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                    metadata["title"] = doc.core_properties.title
                if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                    metadata["author"] = doc.core_properties.author
                if hasattr(doc.core_properties, 'subject') and doc.core_properties.subject:
                    metadata["subject"] = doc.core_properties.subject
                    
        except Exception as e:
            print(f"Failed to extract metadata: {e}")
            
        return metadata
    
    def _convert_with_docx2pdf(self, word_path: str, output_path: str) -> bool:
        """Convert using docx2pdf library (Windows/Mac optimized)"""
        try:
            # docx2pdf works best on Windows with MS Word installed
            docx2pdf_convert(word_path, output_path)
            return os.path.exists(output_path) and os.path.getsize(output_path) > 0
        except Exception as e:
            print(f"docx2pdf conversion failed: {e}")
            return False
    
    def _convert_with_libreoffice(self, word_path: str, output_path: str) -> bool:
        """Convert using LibreOffice headless mode"""
        try:
            # Check if LibreOffice is available
            libreoffice_commands = ['libreoffice', 'soffice']
            libreoffice_cmd = None
            
            for cmd in libreoffice_commands:
                try:
                    subprocess.run([cmd, '--version'], capture_output=True, timeout=5)
                    libreoffice_cmd = cmd
                    break
                except (subprocess.TimeoutExpired, FileNotFoundError):
                    continue
            
            if not libreoffice_cmd:
                print("LibreOffice not found")
                return False
            
            # Create temp directory for conversion
            temp_output_dir = tempfile.mkdtemp()
            
            # Convert using LibreOffice headless
            cmd = [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', temp_output_dir,
                word_path
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0:
                # Find the generated PDF file
                base_name = os.path.splitext(os.path.basename(word_path))[0]
                temp_pdf = os.path.join(temp_output_dir, f"{base_name}.pdf")
                
                if os.path.exists(temp_pdf):
                    shutil.move(temp_pdf, output_path)
                    shutil.rmtree(temp_output_dir, ignore_errors=True)
                    return True
            
            print(f"LibreOffice conversion failed: {result.stderr}")
            shutil.rmtree(temp_output_dir, ignore_errors=True)
            return False
            
        except Exception as e:
            print(f"LibreOffice conversion error: {e}")
            return False
    
    def _convert_with_pandoc(self, word_path: str, output_path: str) -> bool:
        """Convert using Pandoc with enhanced formatting"""
        try:
            # Convert with Pandoc preserving formatting
            pypandoc.convert_file(
                word_path,
                'pdf',
                outputfile=output_path,
                extra_args=[
                    '--pdf-engine=weasyprint',  # Use WeasyPrint for better formatting
                    '--css=styles.css',         # Custom CSS for styling
                    '--standalone',             # Create standalone document
                    '--preserve-tabs',          # Preserve tab formatting
                    '--table-of-contents'       # Include TOC if present
                ]
            )
            return os.path.exists(output_path) and os.path.getsize(output_path) > 0
            
        except Exception as e:
            print(f"Pandoc conversion failed: {e}")
            try:
                # Fallback without extra args
                pypandoc.convert_file(word_path, 'pdf', outputfile=output_path)
                return os.path.exists(output_path) and os.path.getsize(output_path) > 0
            except Exception as e2:
                print(f"Pandoc fallback also failed: {e2}")
                return False
    
    def _convert_with_reportlab_fallback(self, word_path: str, output_path: str) -> bool:
        """Manual conversion using ReportLab as last resort"""
        try:
            if not (Document and canvas):
                return False
            
            # Read Word document
            doc = Document(word_path)
            
            # Create PDF document
            from reportlab.platypus import SimpleDocTemplate
            pdf_doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=inch,
                leftMargin=inch,
                topMargin=inch,
                bottomMargin=inch
            )
            
            # Build content
            story = []
            styles = getSampleStyleSheet()
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Determine style based on paragraph formatting
                    style = styles['Normal']
                    if paragraph.style.name.startswith('Heading'):
                        style = styles['Heading1']
                    
                    para = Paragraph(paragraph.text, style)
                    story.append(para)
                    story.append(Spacer(1, 12))
            
            # Handle tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                
                if table_data:
                    pdf_table = Table(table_data)
                    pdf_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 14),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(pdf_table)
                    story.append(Spacer(1, 12))
            
            # Build PDF
            pdf_doc.build(story)
            return os.path.exists(output_path) and os.path.getsize(output_path) > 0
            
        except Exception as e:
            print(f"ReportLab fallback conversion failed: {e}")
            return False

def convert_word_file(word_path: str, output_dir: str = None) -> Dict[str, Any]:
    """
    Main function to convert Word to PDF
    
    Args:
        word_path: Path to Word file
        output_dir: Directory to save converted file
        
    Returns:
        Conversion result dictionary
    """
    converter = WordToPDFConverter(output_dir)
    return converter.convert_word_to_pdf(word_path)

if __name__ == "__main__":
    # Test the converter
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python word_to_pdf_converter.py <word_file>")
        sys.exit(1)
        
    word_file = sys.argv[1]
    result = convert_word_file(word_file)
    
    if result["success"]:
        print(f"Conversion successful: {result['filename']}")
        print(f"Method: {result['method']}")
        print(f"Output: {result['output_path']}")
    else:
        print(f"Conversion failed: {result['error']}")
