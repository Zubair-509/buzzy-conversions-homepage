#!/usr/bin/env python3
"""
Advanced PDF to Word converter with formatting preservation
"""

import os
import json
import tempfile
import traceback
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, Tuple
import uuid

try:
    from pdf2docx import Converter, parse
    print("pdf2docx successfully imported")
except ImportError as e:
    print(f"Warning: pdf2docx not available - {e}")
    Converter = None
    parse = None

try:
    import PyPDF2
    from PyPDF2 import PdfReader
except ImportError:
    print("Warning: PyPDF2 not available")
    PyPDF2 = None

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Warning: python-docx not available")
    Document = None

try:
    import pdfplumber
except ImportError:
    print("Warning: pdfplumber not available")
    pdfplumber = None

try:
    from PIL import Image
except ImportError:
    print("Warning: Pillow not available")
    Image = None

class PDFToWordConverter:
    """Advanced PDF to Word converter with formatting preservation"""
    
    def __init__(self, output_dir: str = None):
        self.output_dir = output_dir or tempfile.mkdtemp()
        Path(self.output_dir).mkdir(parents=True, exist_ok=True)
        
    def convert_pdf_to_word(self, pdf_path: str, output_filename: str = None) -> Dict[str, Any]:
        """
        Convert PDF to Word with advanced formatting preservation
        
        Args:
            pdf_path: Path to the input PDF file
            output_filename: Desired output filename (optional)
            
        Returns:
            Dictionary with conversion results and metadata
        """
        try:
            # Validate input file
            if not os.path.exists(pdf_path):
                return {"success": False, "error": "PDF file not found"}
                
            if not pdf_path.lower().endswith('.pdf'):
                return {"success": False, "error": "File must be a PDF"}
            
            # Generate output filename
            if not output_filename:
                base_name = Path(pdf_path).stem
                output_filename = f"{base_name}_converted.docx"
            elif not output_filename.endswith('.docx'):
                output_filename += '.docx'
                
            output_path = os.path.join(self.output_dir, output_filename)
            
            # Get PDF metadata
            metadata = self._get_pdf_metadata(pdf_path)
            
            # Attempt advanced conversion with pdf2docx
            if Converter:
                success = self._convert_with_pdf2docx(pdf_path, output_path)
                if success:
                    # Enhance the converted document
                    self._enhance_document(output_path, metadata)
                    return {
                        "success": True,
                        "output_path": output_path,
                        "filename": output_filename,
                        "method": "pdf2docx_advanced",
                        "metadata": metadata,
                        "file_size": os.path.getsize(output_path)
                    }
            
            # Fallback to manual conversion
            if pdfplumber and Document:
                success = self._convert_with_fallback(pdf_path, output_path)
                if success:
                    return {
                        "success": True,
                        "output_path": output_path,
                        "filename": output_filename,
                        "method": "pdfplumber_fallback",
                        "metadata": metadata,
                        "file_size": os.path.getsize(output_path)
                    }
            
            return {"success": False, "error": "No suitable conversion method available"}
            
        except Exception as e:
            error_msg = f"Conversion failed: {str(e)}"
            print(f"Error in convert_pdf_to_word: {error_msg}")
            print(traceback.format_exc())
            return {"success": False, "error": error_msg}
    
    def _convert_with_pdf2docx(self, pdf_path: str, output_path: str) -> bool:
        """Convert using pdf2docx library (most advanced)"""
        try:
            # Use pdf2docx's parse function for better control
            parse(
                pdf_file=pdf_path,
                docx_file=output_path,
                start=0,  # Start from first page
                end=None,  # Convert all pages
                pages=None,  # Convert all pages
                password=None,  # No password
                debug=True
            )
            return True
        except Exception as e:
            print(f"pdf2docx conversion failed: {e}")
            # Try alternative pdf2docx method
            try:
                cv = Converter(pdf_path)
                cv.convert(output_path)
                cv.close()
                return True
            except Exception as e2:
                print(f"Alternative pdf2docx method also failed: {e2}")
                return False
    
    def _convert_with_fallback(self, pdf_path: str, output_path: str) -> bool:
        """Fallback conversion using pdfplumber + python-docx"""
        try:
            doc = Document()
            
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    if page_num > 0:
                        # Add page break between pages
                        doc.add_page_break()
                    
                    # Add page header
                    if page_num == 0:
                        header_para = doc.add_paragraph(f"Converted from PDF - Page {page_num + 1}")
                        header_para.style = 'Heading 3'
                    
                    # Extract text with better formatting
                    text = page.extract_text(layout=True, x_tolerance=3, y_tolerance=3)
                    if text and text.strip():
                        # Split text into paragraphs and add them
                        paragraphs = text.split('\n\n')
                        for para_text in paragraphs:
                            if para_text.strip():
                                para = doc.add_paragraph(para_text.strip())
                                # Basic formatting based on text characteristics
                                if len(para_text.strip()) < 100 and para_text.strip().isupper():
                                    para.style = 'Heading 2'
                                elif para_text.strip().endswith(':'):
                                    para.style = 'Heading 3'
                    
                    # Extract and add tables with better formatting
                    tables = page.extract_tables()
                    if tables:
                        for table_data in tables:
                            if table_data:
                                self._add_table_to_doc(doc, table_data)
                    
                    # Handle images better
                    if hasattr(page, 'images') and page.images:
                        img_para = doc.add_paragraph(f"\n[Note: {len(page.images)} image(s) found on page {page_num + 1} - Image extraction requires advanced processing]\n")
                        img_para.style = 'Intense Quote'
            
            # Add document properties
            doc.core_properties.title = "PDF Conversion Result"
            doc.core_properties.author = "PDF to Word Converter"
            doc.core_properties.comments = "Converted using fallback method with pdfplumber + python-docx"
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Fallback conversion failed: {e}")
            return False
    
    def _add_table_to_doc(self, doc, table_data):
        """Add table to Word document with formatting"""
        if not table_data or not any(table_data):
            return
            
        # Filter out empty rows
        filtered_table = [row for row in table_data if row and any(cell for cell in row if cell)]
        
        if not filtered_table:
            return
            
        # Create table in Word document
        rows = len(filtered_table)
        cols = max(len(row) for row in filtered_table)
        
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        for i, row in enumerate(filtered_table):
            for j, cell in enumerate(row):
                if j < len(table.rows[i].cells):
                    table.rows[i].cells[j].text = str(cell) if cell else ""
    
    def _get_pdf_metadata(self, pdf_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF file"""
        metadata = {"pages": 0, "has_images": False, "has_tables": False}
        
        try:
            if PyPDF2:
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    metadata["pages"] = len(pdf_reader.pages)
                    
                    # Get document info
                    if pdf_reader.metadata:
                        metadata.update({
                            "title": pdf_reader.metadata.get('/Title', ''),
                            "author": pdf_reader.metadata.get('/Author', ''),
                            "subject": pdf_reader.metadata.get('/Subject', ''),
                            "creator": pdf_reader.metadata.get('/Creator', '')
                        })
            
            # Check for images and tables using pdfplumber
            if pdfplumber:
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        if hasattr(page, 'images') and page.images:
                            metadata["has_images"] = True
                        
                        tables = page.extract_tables()
                        if tables:
                            metadata["has_tables"] = True
                            
        except Exception as e:
            print(f"Error extracting metadata: {e}")
            
        return metadata
    
    def _enhance_document(self, doc_path: str, metadata: Dict[str, Any]):
        """Enhance the converted document with additional formatting"""
        try:
            if not Document:
                return
                
            doc = Document(doc_path)
            
            # Add document properties
            doc.core_properties.title = metadata.get("title", "Converted Document")
            doc.core_properties.author = metadata.get("author", "PDF Converter")
            doc.core_properties.comments = f"Converted from PDF - {metadata.get('pages', 0)} pages"
            
            # Save enhanced document
            doc.save(doc_path)
            
        except Exception as e:
            print(f"Error enhancing document: {e}")

def convert_pdf_file(pdf_path: str, output_dir: str = None) -> Dict[str, Any]:
    """
    Main function to convert PDF to Word
    
    Args:
        pdf_path: Path to PDF file
        output_dir: Directory to save converted file
        
    Returns:
        Conversion result dictionary
    """
    converter = PDFToWordConverter(output_dir)
    return converter.convert_pdf_to_word(pdf_path)

if __name__ == "__main__":
    # Test the converter
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python pdf_to_word_converter.py <pdf_file>")
        sys.exit(1)
        
    pdf_file = sys.argv[1]
    result = convert_pdf_file(pdf_file)
    
    if result["success"]:
        print(f"Conversion successful: {result['filename']}")
        print(f"Method: {result['method']}")
        print(f"Output: {result['output_path']}")
    else:
        print(f"Conversion failed: {result['error']}")